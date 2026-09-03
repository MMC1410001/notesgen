"""Render the generated Markdown into .docx files for Google Docs.

Only the Markdown subset the prompts actually produce is supported: headings,
bullet/numbered lists, fenced code, blockquotes, rules, and inline
bold/italic/code. Headings map onto Word heading styles, which is what makes
the Google Docs outline pane work after upload.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

CODE_FONT = "Consolas"
CODE_SHADE = "F2F3F5"
CODE_COLOR = RGBColor(0x1A, 0x1A, 0x1A)
QUOTE_COLOR = RGBColor(0x6A, 0x4B, 0x00)

FENCE = re.compile(r"^\s*```(\w*)\s*$")
HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET = re.compile(r"^(\s*)[-*+]\s+(.*)$")
NUMBERED = re.compile(r"^(\s*)(\d+)[.)]\s+(.*)$")
QUOTE = re.compile(r"^\s*>\s?(.*)$")
RULE = re.compile(r"^\s*([-*_])\1{2,}\s*$")
INLINE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|__[^_]+__|(?<!\w)\*[^*\n]+\*(?!\w))")


def _shade(paragraph, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    paragraph._p.get_or_add_pPr().append(shd)


def _add_inline(paragraph, text: str, *, base_italic: bool = False) -> None:
    """Render **bold**, *italic* and `code` spans into a paragraph."""
    for token in INLINE.split(text):
        if not token:
            continue
        if token.startswith("`") and token.endswith("`") and len(token) > 1:
            run = paragraph.add_run(token[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt(9.5)
            run.font.color.rgb = CODE_COLOR
        elif (token.startswith("**") and token.endswith("**")) or (
            token.startswith("__") and token.endswith("__")
        ):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(token)
        if base_italic:
            run.italic = True


def _normalise_heading_levels(lines: list[str]) -> dict[int, int]:
    """Map the heading levels actually used onto a contiguous 1..n range.

    The lecture notes use `#` for the title then `###` for the four sections,
    which would leave a hole at Heading 2 and make the Docs outline look
    broken. Collapsing the observed levels fixes that without the prompts
    having to know about it.
    """
    used = sorted({len(m.group(1)) for line in lines if (m := HEADING.match(line))})
    return {level: i + 1 for i, level in enumerate(used)}


def _add_code_block(doc, code_lines: list[str]) -> None:
    for line in code_lines or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Pt(18)
        run = p.add_run(line if line.strip() else " ")
        run.font.name = CODE_FONT
        run.font.size = Pt(9.5)
        run.font.color.rgb = CODE_COLOR
        _shade(p, CODE_SHADE)


def render_markdown(doc, markdown: str, level_map: dict[int, int] | None = None) -> None:
    lines = markdown.splitlines()
    level_map = level_map or _normalise_heading_levels(lines)

    i = 0
    while i < len(lines):
        line = lines[i]

        fence = FENCE.match(line)
        if fence:
            block: list[str] = []
            i += 1
            while i < len(lines) and not FENCE.match(lines[i]):
                block.append(lines[i])
                i += 1
            i += 1  # closing fence
            _add_code_block(doc, block)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue

        if not line.strip():
            i += 1
            continue

        if RULE.match(line):
            i += 1
            continue

        m = HEADING.match(line)
        if m:
            level = level_map.get(len(m.group(1)), len(m.group(1)))
            heading = doc.add_heading("", level=min(level, 4))
            _add_inline(heading, m.group(2))
            i += 1
            continue

        m = QUOTE.match(line)
        if m:
            quoted = [m.group(1)]
            i += 1
            while i < len(lines) and (q := QUOTE.match(lines[i])):
                quoted.append(q.group(1))
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_before = Pt(6)
            _add_inline(p, " ".join(x for x in quoted if x.strip()))
            for run in p.runs:
                run.font.color.rgb = QUOTE_COLOR
            _shade(p, "FFF8E5")
            continue

        m = BULLET.match(line)
        if m:
            depth = min(len(m.group(1)) // 2, 2)
            style = "List Bullet" if depth == 0 else f"List Bullet {depth + 1}"
            p = doc.add_paragraph(style=_style_or_default(doc, style, "List Bullet"))
            _add_inline(p, m.group(2))
            i += 1
            continue

        m = NUMBERED.match(line)
        if m:
            depth = min(len(m.group(1)) // 2, 2)
            style = "List Number" if depth == 0 else f"List Number {depth + 1}"
            p = doc.add_paragraph(style=_style_or_default(doc, style, "List Number"))
            _add_inline(p, m.group(3))
            i += 1
            continue

        p = doc.add_paragraph()
        _add_inline(p, line.strip())
        i += 1


def _style_or_default(doc, wanted: str, fallback: str) -> str:
    try:
        doc.styles[wanted]
        return wanted
    except KeyError:
        return fallback


def new_document(title: str, subtitle: str) -> "Document":
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph()
    run = sub.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return doc


def build_section_doc(
    title: str,
    subtitle: str,
    parts: list[str],
    out_path: Path,
) -> Path:
    """Assemble one .docx from a list of Markdown documents."""
    parts = list(parts)
    # The section overview opens with its own H1 repeating the section name,
    # which would show up twice under the document Title.
    if parts:
        first = parts[0].lstrip().splitlines()
        if first and first[0].startswith("# ") and first[0][2:].strip() == title.strip():
            parts[0] = parts[0].lstrip().split("\n", 1)[1] if "\n" in parts[0].lstrip() else ""

    combined = "\n\n".join(parts)
    level_map = _normalise_heading_levels(combined.splitlines())

    doc = new_document(title, subtitle)
    for n, part in enumerate(parts):
        if n:
            doc.add_page_break()
        render_markdown(doc, part, level_map)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
